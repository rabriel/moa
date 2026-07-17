from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "Mall of the North - Find the Bananas User Manual.docx"

TITLE_BLUE = RGBColor(14, 71, 116)
ACCENT_YELLOW = RGBColor(249, 239, 72)
BODY_COLOR = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "E8EEF5"
TABLE_BORDER = "C7D3E0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER)


def style_run(run, *, size=None, bold=None, color=None, font_name="Calibri"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = alignment
    if text:
        paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    style_run(run, size=11, color=BODY_COLOR)


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    style_run(run, size=11, color=BODY_COLOR)


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, TITLE_BLUE),
        ("Heading 2", 13, 14, 7, TITLE_BLUE),
        ("Heading 3", 12, 10, 5, TITLE_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    if "ManualCallout" not in doc.styles:
        style = doc.styles.add_style("ManualCallout", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = TITLE_BLUE
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15

    return doc


def add_title_block(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Mall of the North\nFind the Bananas User Manual")
    style_run(run, size=26, bold=False, color=BODY_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Detailed operating guide for the player journey, QR hunt flow, admin access, "
        "and day-to-day campaign support."
    )
    style_run(run, size=11, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run("Document date: 17 July 2026")
    style_run(run, size=10, color=MUTED)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)

    headers = [("Item", 1.8), ("Details", 4.7)]
    for idx, (label, width) in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        style_run(run, size=11, bold=True, color=TITLE_BLUE)

    rows = [
        ("Campaign name", "Meet Henry & James / Find the Bananas QR clue hunt"),
        ("Primary audience", "Mall visitors entering and playing on their phones"),
        ("Game length", "20 store locations"),
        ("Core player actions", "Register, log in, scan QR codes, follow clues, complete all locations"),
        ("Admin actions", "Sign in, view entries, review progress, export entries to Excel"),
        ("Technology", "Laravel 10, Blade, Bootstrap 5, jQuery, MySQL, Vite"),
    ]

    for label, detail in rows:
        row = table.add_row().cells
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_width(row[0], 1.8)
        set_cell_width(row[1], 4.7)
        row[0].paragraphs[0].add_run(label).bold = True
        row[1].paragraphs[0].add_run(detail)


def add_store_table(doc):
    stores = [
        ("1", "Ackermans", "ackermans", "Intro clue starts here. Scanning Ackermans reveals Lacoste."),
        ("2", "Lacoste", "lacoste", "Scanning reveals the clue for Spec-Savers."),
        ("3", "Spec-Savers", "spec-savers", "Scanning reveals the clue for Baby City."),
        ("4", "Baby City", "baby-city", "Scanning reveals the clue for The Fun Company."),
        ("5", "The Fun Company", "the-fun-company", "Scanning reveals the clue for Expedition North."),
        ("6", "Expedition North", "expedition-north", "Scanning reveals the clue for Coricraft."),
        ("7", "Coricraft", "coricraft", "Scanning reveals the clue for Legends Barbershop."),
        ("8", "Legends Barbershop", "legends-barbershop", "Scanning reveals the clue for Clicks."),
        ("9", "Clicks", "clicks", "Scanning reveals the clue for Lovisa."),
        ("10", "Lovisa", "lovisa", "Scanning reveals the clue for Destinations by Frasers."),
        ("11", "Destinations by Frasers", "destinations-by-frasers", "Scanning reveals the clue for Freedom of Movement."),
        ("12", "Freedom of Movement", "freedom-of-movement", "Scanning reveals the clue for Sorbet."),
        ("13", "Sorbet", "sorbet", "Scanning reveals the clue for Old School."),
        ("14", "Old School", "old-school", "Scanning reveals the clue for Le Creuset."),
        ("15", "Le Creuset", "le-creuset", "Scanning reveals the clue for PNA."),
        ("16", "PNA", "pna", "Scanning reveals the clue for Crocs."),
        ("17", "Crocs", "crocs", "Scanning reveals the clue for Cell C."),
        ("18", "Cell C", "cell-c", "Scanning reveals the clue for Totalsports."),
        ("19", "Totalsports", "totalsports", "Scanning reveals the clue for Spur."),
        ("20", "Spur", "spur", "Final stop. Completion screen follows when all required scans are logged."),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)

    headers = [("No.", 0.45), ("Store", 1.7), ("Slug", 1.45), ("What happens", 2.9)]
    for idx, (label, width) in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        style_run(run, size=10, bold=True, color=TITLE_BLUE)

    for record in stores:
        row = table.add_row().cells
        for idx, value in enumerate(record):
            widths = [0.45, 1.7, 1.45, 2.9]
            set_cell_width(row[idx], widths[idx])
            row[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            paragraph = row[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            style_run(run, size=10.5, color=BODY_COLOR)


def add_qr_examples_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)

    headers = [("Example", 2.4), ("Purpose", 4.1)]
    for idx, (label, width) in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_FILL)
        run = cell.paragraphs[0].add_run(label)
        style_run(run, size=11, bold=True, color=TITLE_BLUE)

    rows = [
        ("https://findthebananas.co.za/scan/ackermans", "Ackermans intro / first active QR destination"),
        ("https://findthebananas.co.za/scan/lacoste", "Second store in the sequence"),
        ("https://findthebananas.co.za/scan/spec-savers", "Third store in the sequence"),
        ("https://findthebananas.co.za/scan/spur", "Last store in the sequence"),
    ]

    for example, purpose in rows:
        row = table.add_row().cells
        set_cell_width(row[0], 2.4)
        set_cell_width(row[1], 4.1)
        row[0].paragraphs[0].add_run(example)
        row[1].paragraphs[0].add_run(purpose)


def build_document():
    doc = configure_document()
    add_title_block(doc)

    doc.add_heading("1. Purpose of this manual", level=1)
    doc.add_paragraph(
        "This manual explains how the Mall of the North Find the Bananas web application works, "
        "how players move through the hunt, what admins can manage, and what support staff should "
        "know when helping visitors on site."
    )
    doc.add_paragraph(
        "It is intended for campaign managers, mall staff, promoters, digital support teams, and "
        "anyone who needs a clear reference for how the application should behave."
    )

    doc.add_heading("2. Quick reference", level=1)
    add_summary_table(doc)

    doc.add_heading("3. How the game works at a high level", level=1)
    add_number(doc, "A visitor lands on the campaign home page.")
    add_number(doc, "The visitor registers with their name, surname, email address, cellphone number, and password.")
    add_number(doc, "After registration, the visitor is logged into the session and redirected into the hunt.")
    add_number(doc, "The hunt starts with the Ackermans intro clue and progress displays as 0/20.")
    add_number(doc, "The player scans the QR code at each correct store location in sequence.")
    add_number(doc, "Each successful scan logs that store against the player and reveals the next clue.")
    add_number(doc, "If a player scans the same store again, the clue is shown again but progress does not increase.")
    add_number(doc, "If a player jumps ahead to the wrong next store, the warning message tells them they missed banana codes.")
    add_number(doc, "After all 20 stores are correctly logged, the completion screen is shown.")

    doc.add_heading("4. Player journey in detail", level=1)
    doc.add_heading("4.1 Landing page", level=2)
    doc.add_paragraph(
        "The landing page introduces the campaign, shows the Henry and James creative, and presents "
        "the primary call to action that takes the visitor into registration."
    )
    add_bullet(doc, "Purpose: explain the game quickly and encourage entry.")
    add_bullet(doc, "Main action: begin registration.")
    add_bullet(doc, "Visual style: mobile-first, campaign artwork, blue gradient background, yellow call-to-action button.")

    doc.add_heading("4.2 Registration page", level=2)
    doc.add_paragraph(
        "The registration page collects the player’s core details before they can participate in the hunt."
    )
    add_bullet(doc, "Required fields: Name, Surname, Email, Cellphone, Password, Password Confirmation.")
    add_bullet(doc, "Validation: empty fields are blocked, emails must be valid, password confirmation must match.")
    add_bullet(doc, "One entry per person is supported through player records and validation rules.")
    add_bullet(doc, "On successful registration, the player record is saved and the player is logged into session.")
    add_bullet(doc, "The player is then redirected to the Ackermans intro state.")

    doc.add_heading("4.3 Login page", level=2)
    doc.add_paragraph(
        "Returning players can use the login page instead of registering again."
    )
    add_bullet(doc, "Login fields: email address and password.")
    add_bullet(doc, "Successful login restores access to the player’s existing progress.")
    add_bullet(doc, "A logout link is available under the Terms & Conditions area on gameplay screens.")

    doc.add_heading("4.4 Forgot password page", level=2)
    doc.add_paragraph(
        "Players who forget their password can use the forgot password feature."
    )
    add_bullet(doc, "The reset flow uses player verification details and then updates the password in the database.")
    add_bullet(doc, "This helps support repeat visitors without forcing a new competition entry.")

    doc.add_heading("4.5 Intro Ackermans state", level=2)
    doc.add_paragraph(
        "Immediately after registration, the player does not receive a completed scan yet. Instead, the system "
        "opens the first clue screen for Ackermans with progress still showing 0/20."
    )
    add_bullet(doc, "This is intentional and prevents the first location from being auto-counted.")
    add_bullet(doc, "The player still needs to physically go to Ackermans and scan its QR code.")
    add_bullet(doc, "The intro screen does not show a success message because nothing has been scanned yet.")

    doc.add_heading("5. QR scanning and clue logic", level=1)
    doc.add_paragraph(
        "The application uses dedicated QR destination links in the format `/scan/{store-slug}`. Each store "
        "has its own unique route and its own place in the overall hunt sequence."
    )
    doc.add_heading("5.1 Important rule", level=2)
    callout = doc.add_paragraph(style="ManualCallout")
    callout.add_run("A QR code should only count when the player scans the correct next store in the hunt order.")

    doc.add_heading("5.2 What happens after a correct scan", level=2)
    add_bullet(doc, "The app checks that the player is logged in.")
    add_bullet(doc, "The store is matched by its slug.")
    add_bullet(doc, "If the player has not logged that store before, a new visit record is created.")
    add_bullet(doc, "Progress increases by one.")
    add_bullet(doc, "A store-specific success message is shown.")
    add_bullet(doc, "The clue card then points to the next store, not the one just scanned.")

    doc.add_heading("5.3 What happens on a duplicate scan", level=2)
    doc.add_paragraph(
        "If the player scans a store that has already been logged, the system does not create a second visit."
    )
    add_bullet(doc, "Progress stays the same.")
    add_bullet(doc, "The player sees: “You've already logged this store - here's the clue again.”")
    add_bullet(doc, "The clue card is still shown so the player can continue the hunt.")

    doc.add_heading("5.4 What happens if a player jumps ahead", level=2)
    doc.add_paragraph(
        "The hunt is sequence-based. If the player scans a store that is not the expected next location, "
        "the application does not log the visit."
    )
    add_bullet(doc, "Progress does not increase.")
    add_bullet(doc, "The player sees: “Boo-doo! You missed some banana codes! Go hunt again!”")
    add_bullet(doc, "The clue card points back to the expected next store so the player can return to the correct path.")

    doc.add_heading("5.5 Completion behaviour", level=2)
    add_bullet(doc, "The hunt target is 20 out of 20 locations.")
    add_bullet(doc, "Once the player has logged all required locations, they are redirected to the completion screen.")
    add_bullet(doc, "The completion view celebrates the finished hunt and shows the winner/final prize copy.")

    doc.add_heading("6. Progress bar behaviour", level=1)
    doc.add_paragraph(
        "The game shows a 20-step progress tracker. Each correct store visit fills one progress slot."
    )
    add_bullet(doc, "Initial Ackermans intro screen: 0/20.")
    add_bullet(doc, "After scanning Ackermans correctly: 1/20.")
    add_bullet(doc, "After scanning Lacoste correctly: 2/20.")
    add_bullet(doc, "The progress bar uses the campaign’s on/off assets and should visually match the mobile designs.")
    add_bullet(doc, "Duplicate or out-of-sequence scans must never increase the bar.")

    doc.add_heading("7. Store order and route reference", level=1)
    doc.add_paragraph(
        "The table below lists the active store sequence currently used by the application."
    )
    add_store_table(doc)

    doc.add_heading("8. QR link examples", level=1)
    doc.add_paragraph(
        "Actual QR artwork can point to the live domain. The examples below show the standard route format."
    )
    add_qr_examples_table(doc)

    doc.add_heading("9. Player-facing copy currently used in the app", level=1)
    add_bullet(doc, "Scan note: “Instructions: Open your device camera to Scan”")
    add_bullet(doc, "Duplicate warning: “You've already logged this store - here's the clue again.”")
    add_bullet(doc, "Missed sequence warning: “Boo-doo! You missed some banana codes! Go hunt again!”")
    add_bullet(doc, "Completion logic: all 20 stores must be logged before the player reaches the final completion view.")
    add_bullet(doc, "Logout display: “FirstName | Logout” shown under the Terms & Conditions area on gameplay screens.")

    doc.add_heading("10. Admin area", level=1)
    doc.add_paragraph(
        "The application includes a restricted admin area for viewing and exporting competition entries."
    )
    add_bullet(doc, "Admins sign in with admin credentials, not player credentials.")
    add_bullet(doc, "The admin entries screen uses the same campaign colour direction as the main app.")
    add_bullet(doc, "The dashboard shows player entries, store counts, visited store data, completion status, and created dates.")
    add_bullet(doc, "Admins can download the entries into Excel from the export action.")
    add_bullet(doc, "Table text should remain readable with black text in the data area.")

    doc.add_heading("11. Terms & Conditions page", level=1)
    doc.add_paragraph(
        "A Terms & Conditions page exists as part of the campaign structure. The visual design already reserves "
        "space for the Terms link on the player journey."
    )
    add_bullet(doc, "The page can be updated with final legal copy when approved.")
    add_bullet(doc, "The user journey already accommodates the link placement.")

    doc.add_heading("12. Common support scenarios", level=1)
    doc.add_heading("12.1 A player says the first location was already counted", level=2)
    add_bullet(doc, "Expected result: it should not be auto-counted.")
    add_bullet(doc, "Ackermans first appears as the intro clue at 0/20.")
    add_bullet(doc, "The player must still scan the Ackermans QR code in person to get 1/20.")

    doc.add_heading("12.2 A player says progress did not increase", level=2)
    add_bullet(doc, "Check whether they scanned a duplicate location.")
    add_bullet(doc, "Check whether they skipped ahead and scanned the wrong next location.")
    add_bullet(doc, "If either of those happened, no new visit should be logged.")

    doc.add_heading("12.3 A player says the camera button is not scanning inside the site", level=2)
    add_bullet(doc, "Current copy instructs the player to use the device camera directly.")
    add_bullet(doc, "The flow is intentionally simplified so the phone camera handles QR detection more reliably.")
    add_bullet(doc, "If needed, the player should open the phone camera app and scan the next QR code from there.")

    doc.add_heading("12.4 A player sees a page expired message", level=2)
    add_bullet(doc, "The app includes protection to avoid raw 419 errors where possible.")
    add_bullet(doc, "If a session has genuinely expired, the player should log in again and continue from saved progress.")

    doc.add_heading("13. Deployment and live environment notes", level=1)
    add_bullet(doc, "Build frontend assets before deployment so `public/build` is present.")
    add_bullet(doc, "Do not leave a `public/hot` file on the live server.")
    add_bullet(doc, "The live application should use HTTPS and the canonical URL should be enforced.")
    add_bullet(doc, "Database migrations should be run before launch when schema changes are introduced.")
    add_bullet(doc, "Store data should remain aligned with the seeded 20-location campaign order.")

    doc.add_heading("14. Recommended operational checklist", level=1)
    add_number(doc, "Confirm the home page loads correctly on mobile.")
    add_number(doc, "Test new player registration.")
    add_number(doc, "Verify login and forgot password flows.")
    add_number(doc, "Open the Ackermans intro screen and confirm progress starts at 0/20.")
    add_number(doc, "Scan Ackermans and confirm Lacoste is the next clue.")
    add_number(doc, "Test one duplicate scan and confirm progress does not change.")
    add_number(doc, "Test one out-of-sequence scan and confirm the Boo-doo warning appears.")
    add_number(doc, "Reach the last store and confirm the completion experience is correct.")
    add_number(doc, "Check admin login and Excel export.")

    doc.add_heading("15. Document owner notes", level=1)
    doc.add_paragraph(
        "This manual reflects the application behaviour present in the codebase on 17 July 2026. "
        "If the campaign copy, store order, scan rules, or admin permissions change, this document "
        "should be updated so operational teams always work from the current version."
    )

    section = doc.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Find the Bananas User Manual")
    style_run(run, size=9, color=MUTED)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
